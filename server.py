import os
import json
import logging
import asyncio
import re
from typing import Dict
from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

DATA_FILE = "data/data.json"
MENU = {
    "Foods": {"Firfir": 150, "Dulet": 250, "Doro Wot": 400, "Pasta": 200, "Shiro": 120, "Tegabino": 180},
    "Soft Drinks": {"Pepsi": 50, "Coca": 50, "Fanta": 50, "Mirinda": 50},
    "Hot Drinks": {"Tea": 30, "Coffee": 40, "Shorba": 80}
}
FLAT_MENU = {item: price for cat in MENU for item, price in MENU[cat].items()}
PAYMENT_METHODS = ["Cash", "Telebirr", "CBE", "BOA", "Amara Bank"]

data_lock = asyncio.Lock()
app = FastAPI(title="Ethiopian Restaurant App")

os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

# --- DATA STORAGE HANDLERS ---
def initialize_data_file():
    os.makedirs("data", exist_ok=True)
    if not os.path.exists(DATA_FILE):
        with open(DATA_FILE, "w") as f:
            json.dump({"last_ticket": 0, "tickets": {}}, f, indent=2)

def load_data() -> dict:
    try:
        with open(DATA_FILE, "r") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {"last_ticket": 0, "tickets": {}}

def save_data(data: dict) -> None:
    try:
        with open(DATA_FILE, "w") as f:
            json.dump(data, f, indent=2)
    except IOError as e:
        logger.error(f"Error saving data: {e}")

initialize_data_file()

REPORTS_DIR = "data/reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

# --- DAILY REPORT (.docx) GENERATOR ---
CATEGORY_COLORS = {
    "Foods":       {"header_bg": "FEF3C7", "header_text": "92400E"},
    "Soft Drinks": {"header_bg": "E0F2FE", "header_text": "075985"},
    "Hot Drinks":  {"header_bg": "FFE4E6", "header_text": "9F1239"},
}
DEFAULT_CATEGORY_COLOR = {"header_bg": "F1F5F9", "header_text": "334155"}

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, shading=None):
    cell.text = ""
    if shading:
        _shade_cell(cell, shading)
    p = cell.paragraphs[0]
    p.alignment = align
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break(WD_BREAK.LINE)
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

def generate_daily_report_docx() -> io.BytesIO:
    db = load_data()
    tickets_dict = db.get("tickets", {})
    tickets_sorted = sorted(tickets_dict.items(), key=lambda x: int(x[0]))

    paid = [t for _, t in tickets_sorted if t.get("paid")]
    pending = [t for _, t in tickets_sorted if not t.get("paid")]
    grand_revenue = sum(t.get("total", 0) for t in paid)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.top_margin = sec.bottom_margin = Inches(0.6)

    # ---- Title Block ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("🍽️ Habesha Smart Menu")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor.from_string("0F172A")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("ESU Restaurant — Daily Operations Report")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string("B45309")

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = gen_p.add_run(f"Generated: {datetime.now().strftime('%A, %d %B %Y — %I:%M %p')}")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("64748B")

    doc.add_paragraph()

    # ---- Quick Stats Strip ----
    stats = [("Total Tickets", str(len(tickets_sorted))), ("Paid", str(len(paid))),
             ("Pending", str(len(pending))), ("Total Collected", f"{grand_revenue} ብር")]
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, (label, value) in enumerate(stats):
        _set_cell(stats_table.rows[0].cells[col], label, bold=True, size=8, color="CBD5E1",
                   align=WD_ALIGN_PARAGRAPH.CENTER, shading="0F172A")
        _set_cell(stats_table.rows[1].cells[col], value, bold=True, size=15, color="34D399",
                   align=WD_ALIGN_PARAGRAPH.CENTER, shading="1E293B")

    doc.add_paragraph()

    # ---- Section 1: Full Ticket Log ----
    h1 = doc.add_paragraph(); r = h1.add_run("📋 Full Ticket Log"); r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("0F172A")

    log_table = doc.add_table(rows=1, cols=4)
    log_table.style = 'Table Grid'
    headers = ["Order #", "Items", "Total", "Payment / Status"]
    for i, htext in enumerate(headers):
        _set_cell(log_table.rows[0].cells[i], htext, bold=True, size=10, color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER, shading="0F172A")

    for idx, (tid, t) in enumerate(tickets_sorted):
        row = log_table.add_row()
        items_text = "\n".join(f"{q}x {name}" for name, q in t.get("items", {}).items())
        is_paid = t.get("paid", False)
        status_text = t.get("method", "Cash") if is_paid else "Pending"
        row_shading = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        _set_cell(row.cells[0], f"#{tid}", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, shading=row_shading)
        _set_cell(row.cells[1], items_text, size=9, shading=row_shading)
        _set_cell(row.cells[2], f"{t.get('total', 0)} ብር", bold=True, size=10, color="059669",
                   align=WD_ALIGN_PARAGRAPH.CENTER, shading=row_shading)
        _set_cell(row.cells[3], status_text, bold=True, size=9,
                   color="B45309" if is_paid else "DC2626", align=WD_ALIGN_PARAGRAPH.CENTER, shading=row_shading)

    if not tickets_sorted:
        row = log_table.add_row()
        _set_cell(row.cells[0], "No tickets recorded yet.", size=9, color="94A3B8")

    doc.add_paragraph()

    # ---- Section 2: Sales Report by Category ----
    h2 = doc.add_paragraph(); r = h2.add_run("📊 Sales Report"); r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("0F172A")

    sales = {item: 0 for item in FLAT_MENU}
    for t in paid:
        for item, qty in t.get("items", {}).items():
            if item in sales:
                sales[item] += qty

    for cat, cat_items in MENU.items():
        theme = CATEGORY_COLORS.get(cat, DEFAULT_CATEGORY_COLOR)
        cat_p = doc.add_paragraph(); r = cat_p.add_run(cat); r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(theme["header_text"])

        item_rows = [{"name": i, "sold": sales[i], "revenue": sales[i] * FLAT_MENU[i]} for i in cat_items]
        item_rows.sort(key=lambda x: x["sold"], reverse=True)
        cat_total = sum(i["revenue"] for i in item_rows)

        cat_table = doc.add_table(rows=1, cols=3)
        cat_table.style = 'Table Grid'
        for i, htext in enumerate(["Item", "Qty Sold", "Revenue"]):
            _set_cell(cat_table.rows[0].cells[i], htext, bold=True, size=9, color=theme["header_text"],
                       align=WD_ALIGN_PARAGRAPH.CENTER, shading=theme["header_bg"])
        for item in item_rows:
            row = cat_table.add_row()
            _set_cell(row.cells[0], item["name"], size=9)
            _set_cell(row.cells[1], f"x{item['sold']}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(row.cells[2], f"{item['revenue']} ብር", bold=True, size=9, color="059669",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        total_row = cat_table.add_row()
        _set_cell(total_row.cells[0], "Total Collected", bold=True, size=9, color="FFFFFF", shading="0F172A")
        total_row.cells[1].merge(total_row.cells[2])
        _set_cell(total_row.cells[1], f"{cat_total} ብር", bold=True, size=11, color="34D399",
                   align=WD_ALIGN_PARAGRAPH.CENTER, shading="0F172A")
        doc.add_paragraph()

    # ---- Section 3: Revenue Summary by Payment Method ----
    h3 = doc.add_paragraph(); r = h3.add_run("💰 Revenue Summary"); r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("0F172A")

    rev_summary = {m: 0 for m in PAYMENT_METHODS}
    for t in paid:
        method = t.get("method", "Cash")
        rev_summary[method] = rev_summary.get(method, 0) + t.get("total", 0)

    rev_table = doc.add_table(rows=1, cols=2)
    rev_table.style = 'Table Grid'
    _set_cell(rev_table.rows[0].cells[0], "Payment Method", bold=True, size=10, color="FFFFFF",
               shading="0F172A")
    _set_cell(rev_table.rows[0].cells[1], "Amount Collected", bold=True, size=10, color="FFFFFF",
               align=WD_ALIGN_PARAGRAPH.CENTER, shading="0F172A")
    for method, amount in rev_summary.items():
        row = rev_table.add_row()
        _set_cell(row.cells[0], f"💳 {method}", size=10)
        _set_cell(row.cells[1], f"{amount} ብር", bold=True, size=10, color="059669",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    total_row = rev_table.add_row()
    _set_cell(total_row.cells[0], "GRAND TOTAL", bold=True, size=12, color="FFFFFF", shading="0F172A")
    _set_cell(total_row.cells[1], f"{grand_revenue} ብር", bold=True, size=14, color="34D399",
               align=WD_ALIGN_PARAGRAPH.CENTER, shading="0F172A")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_report_snapshot() -> dict:
    """Builds the exact same structured data used by the live Smart Menu Report page."""
    db = load_data()
    tickets_dict = db.get("tickets", {})
    tickets_sorted = sorted(tickets_dict.items(), key=lambda x: int(x[0]))

    paid = [t for _, t in tickets_sorted if t.get("paid")]
    pending = [t for _, t in tickets_sorted if not t.get("paid")]
    grand_revenue = sum(t.get("total", 0) for t in paid)

    ticket_log = [{
        "id": tid,
        "items": t.get("items", {}),
        "total": t.get("total", 0),
        "paid": t.get("paid", False),
        "method": t.get("method") if t.get("paid") else None
    } for tid, t in tickets_sorted]

    sales = {item: 0 for item in FLAT_MENU}
    for t in paid:
        for item, qty in t.get("items", {}).items():
            if item in sales:
                sales[item] += qty

    categories = {}
    for cat, cat_items in MENU.items():
        item_rows = [{"name": i, "sold": sales[i], "revenue": sales[i] * FLAT_MENU[i]} for i in cat_items]
        item_rows.sort(key=lambda x: x["sold"], reverse=True)
        categories[cat] = {"items": item_rows, "category_total": sum(i["revenue"] for i in item_rows)}

    rev_summary = {m: 0 for m in PAYMENT_METHODS}
    for t in paid:
        method = t.get("method", "Cash")
        rev_summary[method] = rev_summary.get(method, 0) + t.get("total", 0)

    return {
        "stats": {"total": len(tickets_sorted), "paid": len(paid), "pending": len(pending), "grand_total": grand_revenue},
        "ticket_log": ticket_log,
        "categories": categories,
        "revenue_breakdown": rev_summary,
        "generated_at": datetime.now().strftime('%A, %d %B %Y — %I:%M %p')
    }

def save_report_to_archive(reason="manual") -> str:
    """Generates the report .docx, saves it into the archive folder, and also
    saves a JSON snapshot of the exact data used so it can be re-viewed later
    exactly as it was at that moment. Returns the docx filename."""
    snapshot = build_report_snapshot()
    buf = generate_daily_report_docx()

    base = f"Habesha_Report_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}_{reason}"
    docx_filename = base + ".docx"
    json_filename = base + ".json"

    with open(os.path.join(REPORTS_DIR, docx_filename), "wb") as f:
        f.write(buf.read())
    with open(os.path.join(REPORTS_DIR, json_filename), "w") as f:
        json.dump(snapshot, f, indent=2)

    return docx_filename

def parse_report_docx(path: str) -> dict:
    """Reads a previously saved report .docx and reconstructs the structured
    data (stats, ticket log, sales by category, revenue breakdown) straight
    from its tables, so a saved snapshot can be rendered exactly as it was
    at save time, not just with today's live data."""
    doc = Document(path)
    tables = doc.tables

    def cell_text(cell):
        return cell.text.strip()

    generated_at = ""
    for p in doc.paragraphs:
        if p.text.strip().startswith("Generated:"):
            generated_at = p.text.strip().replace("Generated:", "").strip()
            break

    stats = {"total": 0, "paid": 0, "pending": 0, "grand_total": 0}
    ticket_log = []
    categories = {}
    revenue_breakdown = {}

    if len(tables) >= 1:
        stats_table = tables[0]
        labels = [cell_text(c) for c in stats_table.rows[0].cells]
        values = [cell_text(c) for c in stats_table.rows[1].cells]
        for label, value in zip(labels, values):
            num = re.sub(r"[^\d]", "", value)
            num = int(num) if num else 0
            if label == "Total Tickets":
                stats["total"] = num
            elif label == "Paid":
                stats["paid"] = num
            elif label == "Pending":
                stats["pending"] = num
            elif label == "Total Collected":
                stats["grand_total"] = num

    if len(tables) >= 2:
        log_table = tables[1]
        for row in log_table.rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) < 4 or not cells[0].startswith("#"):
                continue
            tid = cells[0].lstrip("#")
            items = {}
            for line in cells[1].split("\n"):
                m = re.match(r"(\d+)x\s+(.*)", line.strip())
                if m:
                    items[m.group(2)] = int(m.group(1))
            total = int(re.sub(r"[^\d]", "", cells[2]) or 0)
            status = cells[3]
            paid = status != "Pending"
            ticket_log.append({
                "id": tid, "items": items, "total": total,
                "paid": paid, "method": status if paid else None
            })

    cat_names = list(MENU.keys())
    table_idx = 2
    for cat in cat_names:
        if table_idx >= len(tables):
            break
        cat_table = tables[table_idx]
        item_rows = []
        for row in cat_table.rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) < 3 or cells[0] == "Total Collected":
                continue
            name = cells[0]
            sold = int(re.sub(r"[^\d]", "", cells[1]) or 0)
            revenue = int(re.sub(r"[^\d]", "", cells[2]) or 0)
            item_rows.append({"name": name, "sold": sold, "revenue": revenue})
        category_total = sum(i["revenue"] for i in item_rows)
        categories[cat] = {"items": item_rows, "category_total": category_total}
        table_idx += 1

    if table_idx < len(tables):
        rev_table = tables[table_idx]
        for row in rev_table.rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) < 2:
                continue
            method = cells[0].replace("💳", "").strip()
            amount = int(re.sub(r"[^\d]", "", cells[1]) or 0)
            if method == "GRAND TOTAL":
                continue
            revenue_breakdown[method] = amount

    return {
        "stats": stats,
        "ticket_log": ticket_log,
        "categories": categories,
        "revenue_breakdown": revenue_breakdown,
        "generated_at": generated_at,
    }

class OrderSubmit(BaseModel):
    items: Dict[str, int]

class PaymentSubmit(BaseModel):
    method: str

# --- SYSTEM NAVIGATION: FIXED MENU BUTTON + SLIDE-OUT PANEL ---
def inject_navigation(active_page: str, html_content: str) -> str:
    """
    Injects a menu button fixed to the top-left of the viewport (position: fixed,
    so it never scrolls away on any page or screen size, including mobile).
    Tapping it slides out a full-height nav panel with links to every page.
    The KDS ticket monitor keeps its clean read-only header, just shifted to
    make room for the button, and still gets the same fixed menu button/panel.
    """
    active_btn = "flex items-center gap-3 px-4 py-3.5 rounded-xl font-black text-sm transition-all bg-amber-600 text-white shadow-md shadow-amber-600/20 border border-amber-500 no-underline"
    idle_btn = "flex items-center gap-3 px-4 py-3.5 rounded-xl font-bold text-sm text-slate-600 border border-transparent hover:bg-slate-100 hover:text-slate-900 transition-all no-underline"

    nav_links_html = f"""
            <nav class="flex flex-col gap-1.5">
                <a href="/" class="{active_btn if active_page == 'customer' else idle_btn}">
                    <span class="text-xl">🛒</span>
                    <span>Customer App</span>
                </a>
                <a href="/admin/tickets" class="{active_btn if active_page == 'tickets' else idle_btn}">
                    <span class="text-xl">📋</span>
                    <span>Active Tickets</span>
                </a>
                <a href="/admin/waiter" class="{active_btn if active_page == 'waiter' else idle_btn}">
                    <span class="text-xl">🧾</span>
                    <span>Waiter Panel</span>
                </a>
                <a href="/admin/sales" class="{active_btn if active_page == 'sales' else idle_btn}">
                    <span class="text-xl">📊</span>
                    <span>Sales Report</span>
                </a>
                <a href="/admin/revenue" class="{active_btn if active_page == 'revenue' else idle_btn}">
                    <span class="text-xl">💰</span>
                    <span>Revenue Summary</span>
                </a>
                <a href="/admin/report" class="{active_btn if active_page == 'report' else idle_btn}">
                    <span class="text-xl">📄</span>
                    <span>Smart Menu Report</span>
                </a>
                <a href="/admin/files" class="{active_btn if active_page == 'files' else idle_btn}">
                    <span class="text-xl">🗂️</span>
                    <span>Files</span>
                </a>
            </nav>
    """

    # Button + overlay + slide-out panel: all position:fixed, so they stay put
    # on screen through any amount of scrolling, on every page.
    fixed_menu_block = f"""
    <!-- Fixed Menu Button: always pinned top-left, survives scrolling -->
    <button onclick="toggleMainMenu()" id="menu-toggle-btn"
        class="fixed top-4 left-4 z-[100] w-12 h-12 flex items-center justify-center bg-slate-900 hover:bg-slate-800 active:scale-90 text-white rounded-xl shadow-lg border border-slate-700 transition-all cursor-pointer">
        <span class="text-xl">☰</span>
    </button>

    <!-- Dim backdrop, click to close -->
    <div id="menu-overlay" onclick="toggleMainMenu()"
        class="fixed inset-0 bg-black/50 z-[90] opacity-0 pointer-events-none transition-opacity duration-300"></div>

    <!-- Slide-out Nav Panel (fixed to viewport, not the page, so it never scrolls) -->
    <aside id="main-sidebar"
        class="fixed top-0 left-0 h-screen w-72 bg-slate-50 border-r border-slate-200 p-5 flex flex-col justify-between z-[95] shadow-2xl -translate-x-full transition-transform duration-300 ease-out overflow-y-auto">
        <div class="space-y-6">
            <div class="flex items-center justify-between border-b border-slate-200/80 pb-4">
                <div class="flex items-center gap-3">
                    <span class="text-3xl">🍽️</span>
                    <div>
                        <h1 class="text-base font-black tracking-tight text-slate-900">የምግብ ቤት ማዘዣ</h1>
                        <p class="text-[10px] text-amber-600 font-black uppercase tracking-wider">Habesha Smart Menu</p>
                    </div>
                </div>
                <button onclick="toggleMainMenu()" class="text-slate-400 hover:text-red-500 text-2xl font-bold leading-none cursor-pointer">✕</button>
            </div>
            {nav_links_html}
        </div>

        <div class="pt-4 border-t border-slate-200/60 text-center">
            <p class="text-[10px] text-slate-400 font-bold uppercase tracking-widest">Management Panel v2.1</p>
        </div>
    </aside>

    <script>
        function toggleMainMenu() {{
            const sb = document.getElementById('main-sidebar');
            const ov = document.getElementById('menu-overlay');
            const isOpen = !sb.classList.contains('-translate-x-full');
            if (isOpen) {{
                sb.classList.add('-translate-x-full');
                ov.classList.add('opacity-0', 'pointer-events-none');
            }} else {{
                sb.classList.remove('-translate-x-full');
                ov.classList.remove('opacity-0', 'pointer-events-none');
            }}
        }}
        document.addEventListener('keydown', function(e) {{
            if (e.key === 'Escape') {{
                const sb = document.getElementById('main-sidebar');
                if (sb && !sb.classList.contains('-translate-x-full')) toggleMainMenu();
            }}
        }});
    </script>
    """

    if active_page == 'tickets':
        # Clean read-only KDS header, shifted right so it doesn't collide with the fixed button
        top_bar = """
        <header class="bg-white border-b-2 border-slate-200 p-4 pl-20 flex items-center justify-between shadow-sm">
            <div class="flex items-center gap-4">
                <span class="text-3xl">👨‍🍳</span>
                <div>
                    <h1 class="text-xl font-black tracking-tight text-slate-900">የወጥ ቤት ማሳያ ሰሌዳ (Kitchen KDS Monitor)</h1>
                    <p class="text-xs text-amber-600 font-black uppercase tracking-widest">Live Cooking Queue • Read Only Mode</p>
                </div>
            </div>
            <div class="flex items-center gap-2 bg-slate-50 px-4 py-2 rounded-xl border-2 border-slate-200">
                <span class="w-2.5 h-2.5 rounded-full bg-emerald-500 animate-pulse"></span>
                <span class="text-xs font-black text-slate-600 uppercase tracking-wider">Live System Sync</span>
            </div>
        </header>
        """
        wrapper_open = '<div class="w-full flex flex-col min-h-screen">'
    else:
        # Non-KDS pages: no in-flow header, just top padding so content clears the fixed button
        top_bar = ""
        wrapper_open = '<div class="w-full min-h-screen pt-20 md:pt-24">'

    modified = html_content.replace("<!-- START_WRAPPER_PLACEHOLDER -->", fixed_menu_block + wrapper_open)
    modified = modified.replace("<!-- NAV_SIDEBAR_PLACEHOLDER -->", top_bar)
    return modified.replace("<!-- END_WRAPPER_PLACEHOLDER -->", '</div>')

# ==================== WEB PAGES ====================

@app.get("/", response_class=HTMLResponse)
def customer_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
        <title>Place Order</title>
        <link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->
        
        <!-- Right Main View Canvas Block Content Window -->
        <main class="flex-1 max-w-6xl px-6 py-6 grid grid-cols-1 lg:grid-cols-3 gap-6">
            
            <div class="lg:col-span-2 space-y-6">
                <div id="category-bar" class="grid grid-cols-3 gap-2 bg-slate-200/80 p-1.5 rounded-2xl border border-slate-300/40"></div>
                <div id="menu-items-grid" class="grid grid-cols-1 sm:grid-cols-2 gap-4"></div>
            </div>
            
            <div class="bg-white border-2 border-slate-200 rounded-2xl p-5 shadow-sm h-fit lg:sticky lg:top-6">
                <div class="flex justify-between items-center border-b border-slate-100 pb-3 mb-4">
                    <h2 class="font-black text-slate-900 flex items-center gap-2 text-base">🛒 የትዕዛዝ ዝርዝር (Cart)</h2>
                    <button onclick="clearCart()" class="text-xs font-black text-red-500 py-1.5 px-3 bg-red-50 hover:bg-red-100 rounded-xl cursor-pointer transition-all">Clear</button>
                </div>
                <div id="cart-list" class="space-y-2 max-h-64 overflow-y-auto mb-4 pr-1"></div>
                <div class="border-t border-slate-100 pt-4 space-y-3">
                    <div class="flex justify-between items-baseline font-black text-slate-900">
                        <span class="text-sm text-slate-500">ጠቅላላ ድምር:</span>
                        <span id="cart-total" class="text-2xl text-emerald-600 font-black">0 ብር</span>
                    </div>
                    <button onclick="commitOrder()" class="w-full bg-slate-900 hover:bg-slate-800 active:scale-95 text-white font-black py-4 rounded-xl shadow-md cursor-pointer transition-all text-sm uppercase tracking-wide">Confirm & Send Order</button>
                </div>
            </div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->
        
        <script>
            let menuCache = {}, currentCategory = "Foods", cart = {};
            
            async function start() {
                try {
                    menuCache = await (await fetch('/api/menu')).json();
                    renderTabs(); 
                    renderProducts(); 
                    syncCart();
                } catch (err) {
                    console.error("Initialization error: ", err);
                }
            }
            
            function renderTabs() {
                document.getElementById('category-bar').innerHTML = Object.keys(menuCache).map(cat => `
                    <button onclick="currentCategory='${cat}'; renderTabs(); renderProducts();" class="py-3 px-2 text-xs font-black uppercase rounded-xl transition-all cursor-pointer text-center ${currentCategory === cat ? 'bg-white text-slate-900 shadow-sm font-black border border-slate-200' : 'text-slate-500 hover:text-slate-900'}">${cat}</button>
                `).join('');
            }
            
            function renderProducts() {
                const grid = document.getElementById('menu-items-grid');
                grid.innerHTML = Object.entries(menuCache[currentCategory] || {}).map(([name, price]) => `
                    <div class="bg-white border-2 border-slate-200 p-4 rounded-2xl flex justify-between items-center shadow-sm hover:border-slate-300 transition-all">
                        <div>
                            <h4 class="font-black text-slate-900 text-base">${name}</h4>
                            <p class="text-sm font-bold text-amber-600 mt-0.5">${price} ብር</p>
                        </div>
                        <div class="flex items-center gap-1 bg-slate-100 p-1 rounded-xl border border-slate-200/60">
                            <button onclick="updateQty('${name}', -1)" class="w-11 h-11 flex items-center justify-center font-bold bg-white rounded-lg border border-slate-200 shadow-sm active:scale-90 text-sm cursor-pointer">➖</button>
                            <span class="w-8 text-center text-sm font-black text-slate-800">${cart[name] || 0}</span>
                            <button onclick="updateQty('${name}', 1)" class="w-11 h-11 flex items-center justify-center font-bold bg-white rounded-lg border border-slate-200 shadow-sm active:scale-90 text-sm cursor-pointer">➕</button>
                        </div>
                    </div>
                `).join('');
            }
            
            function updateQty(item, mod) {
                cart[item] = (cart[item] || 0) + mod;
                if(cart[item] <= 0) delete cart[item];
                renderProducts(); syncCart();
            }
            
            function clearCart() { cart = {}; renderProducts(); syncCart(); }
            
            function syncCart() {
                const list = document.getElementById('cart-list');
                let total = 0;
                const entries = Object.entries(cart);
                if(entries.length === 0) {
                    list.innerHTML = '<p class="text-xs text-slate-400 text-center py-6 font-medium">ካርቱ ባዶ ነው (Cart is empty)</p>';
                    document.getElementById('cart-total').innerText = '0 ብር';
                    return;
                }
                const flat = {}; Object.values(menuCache).forEach(c => Object.assign(flat, c));
                list.innerHTML = entries.map(([name, count]) => {
                    total += flat[name] * count;
                    return `<div class="flex justify-between items-center bg-slate-50 border border-slate-100 p-3 rounded-xl text-xs font-medium"><span class="text-slate-700"><b class="bg-slate-200 text-slate-800 px-1.5 py-0.5 rounded mr-1 font-bold">${count}x</b> ${name}</span><span class="font-bold text-slate-900">${flat[name] * count} ብር</span></div>`;
                }).join('');
                document.getElementById('cart-total').innerText = total + ' ብር';
            }
            
            async function commitOrder() {
                if(Object.keys(cart).length === 0) return alert("ካርቱ ባዶ ነው! Please add items.");
                const res = await fetch('/api/orders', { method: 'POST', headers: {'Content-Type': 'application/json'}, body: JSON.stringify({ items: cart }) });
                if(res.ok) { alert("🎉 ትዕዛዝ ተልኳል! Order Placed Successfully."); clearCart(); }
            }
            
            start();
        </script>
    </body>
    </html>
    """
    return inject_navigation('customer', template)

@app.get("/admin/tickets", response_class=HTMLResponse)
def admin_tickets_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Kitchen Display System</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->
        
        <main class="flex-1 max-w-full px-6 py-6">
            <!-- Order count header -->
            <div class="flex items-center justify-between mb-5">
                <h2 class="text-lg font-black text-slate-900 uppercase tracking-widest">Active Orders</h2>
                <div class="flex items-center gap-2 bg-white px-4 py-2 rounded-xl border-2 border-slate-200 shadow-sm">
                    <span class="w-2 h-2 rounded-full bg-emerald-500 animate-pulse"></span>
                    <span id="order-count-label" class="text-xs font-black text-slate-600 uppercase tracking-wider">0 Orders</span>
                </div>
            </div>

            <!-- All tickets rendered at once, plain wrapping grid, fixed card size -->
            <div id="tickets-list" class="grid grid-cols-1 md:grid-cols-2 xl:grid-cols-3 gap-6"></div>

            <!-- Overflow banner: shown only when there are more orders than fit -->
            <div id="overflow-banner" class="hidden mt-6 text-center py-3 px-4 bg-red-50 border-2 border-red-200 rounded-xl">
                <span class="font-black text-red-500 uppercase tracking-wide text-sm animate-pulse">⏳ More Orders Are There — <span id="overflow-count">0</span> waiting</span>
            </div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->

        <script>
            let globalTickets = [];

            async function syncTicketsData() {
                try {
                    const response = await fetch('/api/admin/tickets');
                    const allTickets = await response.json();
                    // Kitchen ONLY cares about working on UNPAID active tickets.
                    globalTickets = allTickets.filter(t => !t.paid);
                    renderAllTickets();
                } catch (err) {
                    console.error("Data tracking failure: ", err);
                }
            }

            function renderAllTickets() {
                const panel = document.getElementById('tickets-list');
                const overflowBanner = document.getElementById('overflow-banner');
                const overflowCount = document.getElementById('overflow-count');
                const countLabel = document.getElementById('order-count-label');

                countLabel.innerText = globalTickets.length + (globalTickets.length === 1 ? ' Order' : ' Orders');

                if (globalTickets.length === 0) {
                    panel.innerHTML = `
                        <div class="col-span-full text-center py-20 bg-white rounded-2xl border-2 border-dashed border-slate-200 text-slate-400 shadow-sm">
                            <span class="text-5xl block mb-4">✅</span>
                            <h3 class="text-xl font-black text-slate-700">ሁሉም ትዕዛዞች ተጠናቀዋል!</h3>
                            <p class="text-sm text-slate-400 mt-1">No pending active orders on the kitchen line.</p>
                        </div>
                    `;
                    overflowBanner.classList.add('hidden');
                    return;
                }

                const MAX_VISIBLE = 9;
                const visible = globalTickets.slice(0, MAX_VISIBLE);
                const hiddenCount = globalTickets.length - visible.length;

                panel.innerHTML = visible.map(t => {
                    const items = Object.entries(t.items).map(([n, q]) => `
                        <div class="flex justify-between items-center py-3 border-b border-slate-100 last:border-0">
                            <span class="font-black text-slate-800 text-lg">${n}</span>
                            <span class="font-black bg-amber-50 text-amber-700 border border-amber-200 px-3 py-1 rounded-full text-base">x${q}</span>
                        </div>
                    `).join('');

                    return `
                        <div class="bg-white border-2 border-slate-200 rounded-2xl shadow-sm overflow-hidden">
                            <div class="flex justify-between items-center px-5 py-3 bg-slate-900">
                                <h4 class="font-black text-amber-400 text-lg tracking-tight">Order #${t.id}</h4>
                                <span class="font-black uppercase tracking-widest text-[10px] px-2 py-1 bg-red-500/10 text-red-400 border border-red-500/30 rounded-lg animate-pulse">In Progress</span>
                            </div>
                            <div class="px-5 py-2 divide-y divide-slate-100">${items}</div>
                        </div>
                    `;
                }).join('');

                if (hiddenCount > 0) {
                    overflowCount.innerText = hiddenCount;
                    overflowBanner.classList.remove('hidden');
                } else {
                    overflowBanner.classList.add('hidden');
                }
            }

            // Initialize App Routines — no rotation, no auto-scroll, no shrinking; plain live-refreshing grid
            syncTicketsData();
            setInterval(syncTicketsData, 4000); // Check DB status updates every 4 seconds in the background
        </script>
    </body>
    </html>
    """
    return inject_navigation('tickets', template)

@app.get("/admin/waiter", response_class=HTMLResponse)
def admin_waiter_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Waiter Panel</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->
        
        <main class="flex-1 max-w-full px-6 py-6">
            <div class="flex items-center justify-between mb-5">
                <h2 class="text-lg font-black text-slate-900 uppercase tracking-widest">Waiter Panel</h2>
                <div class="flex items-center gap-2 bg-white px-4 py-2 rounded-xl border-2 border-slate-200 shadow-sm">
                    <span class="w-2 h-2 rounded-full bg-emerald-500 animate-pulse"></span>
                    <span id="order-count-label" class="text-xs font-black text-slate-600 uppercase tracking-wider">0 Orders</span>
                </div>
            </div>

            <div id="tickets-list" class="grid grid-cols-1 md:grid-cols-2 xl:grid-cols-3 gap-6"></div>

            <div id="overflow-banner" class="hidden mt-6 text-center py-3 px-4 bg-red-50 border-2 border-red-200 rounded-xl">
                <span class="font-black text-red-500 uppercase tracking-wide text-sm animate-pulse">⏳ More Orders Are There — <span id="overflow-count">0</span> waiting</span>
            </div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->

        <script>
            let globalTickets = [];
            let selectedMethods = {}; // ticketId -> chosen payment method, survives re-renders
            const PAYMENT_METHODS = ["Cash", "Telebirr", "CBE", "BOA", "Amara Bank"];

            async function syncTicketsData() {
                try {
                    const response = await fetch('/api/admin/tickets');
                    const allTickets = await response.json();
                    globalTickets = allTickets.filter(t => !t.paid);
                    renderAllTickets();
                } catch (err) {
                    console.error("Data tracking failure: ", err);
                }
            }

            function renderAllTickets() {
                const panel = document.getElementById('tickets-list');
                const overflowBanner = document.getElementById('overflow-banner');
                const overflowCount = document.getElementById('overflow-count');
                const countLabel = document.getElementById('order-count-label');

                countLabel.innerText = globalTickets.length + (globalTickets.length === 1 ? ' Order' : ' Orders');

                if (globalTickets.length === 0) {
                    panel.innerHTML = `
                        <div class="col-span-full text-center py-20 bg-white rounded-2xl border-2 border-dashed border-slate-200 text-slate-400 shadow-sm">
                            <span class="text-5xl block mb-4">✅</span>
                            <h3 class="text-xl font-black text-slate-700">ሁሉም ትዕዛዞች ተጠናቀዋል!</h3>
                            <p class="text-sm text-slate-400 mt-1">No pending active orders on the kitchen line.</p>
                        </div>
                    `;
                    overflowBanner.classList.add('hidden');
                    return;
                }

                const MAX_VISIBLE = 9;
                const visible = globalTickets.slice(0, MAX_VISIBLE);
                const hiddenCount = globalTickets.length - visible.length;

                panel.innerHTML = visible.map(t => {
                    const items = Object.entries(t.items).map(([n, q]) => `
                        <div class="flex justify-between items-center py-3 border-b border-slate-100 last:border-0">
                            <span class="font-black text-slate-800 text-lg">${n}</span>
                            <span class="font-black bg-amber-50 text-amber-700 border border-amber-200 px-3 py-1 rounded-full text-base">x${q}</span>
                        </div>
                    `).join('');

                    const currentSelection = selectedMethods[t.id] || "Cash";
                    const options = PAYMENT_METHODS.map(m => `<option value="${m}" ${m === currentSelection ? 'selected' : ''}>${m}</option>`).join('');

                    return `
                        <div class="bg-white border-2 border-slate-200 rounded-2xl shadow-sm overflow-hidden">
                            <div class="flex justify-between items-center px-5 py-3 bg-slate-900">
                                <h4 class="font-black text-amber-400 text-lg tracking-tight">Order #${t.id}</h4>
                                <span class="font-black uppercase tracking-widest text-[10px] px-2 py-1 bg-red-500/10 text-red-400 border border-red-500/30 rounded-lg animate-pulse">In Progress</span>
                            </div>
                            <div class="px-5 py-2 divide-y divide-slate-100">${items}</div>
                            <div class="px-5 py-4 border-t border-slate-100 bg-slate-50 space-y-3">
                                <div class="flex justify-between items-center">
                                    <span class="text-xs font-black text-slate-500 uppercase tracking-wider">Total</span>
                                    <span class="font-black text-emerald-600 text-lg">${t.total} ብር</span>
                                </div>
                                <select id="pay-method-${t.id}" onchange="selectedMethods['${t.id}'] = this.value" class="w-full text-sm font-bold border-2 border-slate-200 rounded-xl px-3 py-2.5 bg-white cursor-pointer">
                                    ${options}
                                </select>
                                <button onclick="markServed('${t.id}')" class="w-full bg-emerald-600 hover:bg-emerald-700 active:scale-95 text-white font-black py-3 rounded-xl shadow-sm cursor-pointer transition-all text-sm uppercase tracking-wide">✅ Served & Paid</button>
                            </div>
                        </div>
                    `;
                }).join('');

                if (hiddenCount > 0) {
                    overflowCount.innerText = hiddenCount;
                    overflowBanner.classList.remove('hidden');
                } else {
                    overflowBanner.classList.add('hidden');
                }
            }

            async function markServed(ticketId) {
                const select = document.getElementById('pay-method-' + ticketId);
                const method = select.value;
                try {
                    const res = await fetch('/api/admin/tickets/' + ticketId + '/pay', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ method: method })
                    });
                    if (res.ok) {
                        delete selectedMethods[ticketId];
                        syncTicketsData();
                    } else {
                        alert("Could not update ticket. Please try again.");
                    }
                } catch (err) {
                    console.error("Payment update failed: ", err);
                    alert("Network error updating ticket.");
                }
            }

            syncTicketsData();
            setInterval(syncTicketsData, 4000);
        </script>
    </body>
    </html>
    """
    return inject_navigation('waiter', template)

@app.get("/admin/sales", response_class=HTMLResponse)
def admin_sales_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Sales Report</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->
        
        <main class="flex-1 max-w-6xl px-6 py-6 space-y-6">
            <div>
                <h1 class="text-2xl font-black text-slate-900">የሽያጭ ሪፖርት</h1>
                <p class="text-sm font-bold text-slate-400 uppercase tracking-widest">Product Sales Report</p>
            </div>

            <!-- Grand total summary strip -->
            <div id="summary-strip" class="grid grid-cols-1 sm:grid-cols-2 gap-4"></div>

            <!-- Category cards with large item rows -->
            <div id="sales-container" class="grid grid-cols-1 lg:grid-cols-2 gap-6"></div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->
        
        <script>
            const CATEGORY_THEME = {
                "Foods":       { icon: "🍛", accent: "amber",   bg: "bg-amber-50",   text: "text-amber-700",   border: "border-amber-200" },
                "Soft Drinks": { icon: "🥤", accent: "sky",     bg: "bg-sky-50",     text: "text-sky-700",     border: "border-sky-200" },
                "Hot Drinks":  { icon: "☕", accent: "rose",    bg: "bg-rose-50",    text: "text-rose-700",    border: "border-rose-200" }
            };
            const DEFAULT_THEME = { icon: "🍽️", accent: "slate", bg: "bg-slate-50", text: "text-slate-700", border: "border-slate-200" };

            async function loadSales() {
                const data = await (await fetch('/api/admin/reports/sales')).json();

                // --- Grand total summary strip ---
                document.getElementById('summary-strip').innerHTML = `
                    <div class="bg-slate-900 text-white rounded-2xl p-6 shadow-md border border-slate-800">
                        <span class="text-[11px] uppercase tracking-widest font-black text-slate-400 block mb-1">Total Birr Collected</span>
                        <h2 class="text-4xl font-black text-emerald-400 tracking-tight">${data.grand_total_revenue} ብር</h2>
                    </div>
                    <div class="bg-white border-2 border-slate-200 rounded-2xl p-6 shadow-sm">
                        <span class="text-[11px] uppercase tracking-widest font-black text-slate-400 block mb-1">Total Items Sold</span>
                        <h2 class="text-4xl font-black text-slate-900 tracking-tight">${data.grand_total_sold}</h2>
                    </div>
                `;

                // --- Category cards ---
                document.getElementById('sales-container').innerHTML = Object.entries(data.categories).map(([cat, catData]) => {
                    const theme = CATEGORY_THEME[cat] || DEFAULT_THEME;

                    const rows = catData.items.map(i => `
                        <div class="flex justify-between items-center py-4 border-b border-slate-100 last:border-0">
                            <span class="text-slate-800 font-black text-lg sm:text-xl">${i.name}</span>
                            <div class="flex items-center gap-3">
                                <span class="font-black ${theme.bg} ${theme.text} ${theme.border} border px-3 py-1 rounded-full text-base">x${i.sold}</span>
                                <span class="font-black text-emerald-600 text-lg sm:text-xl min-w-[90px] text-right">${i.revenue} ብር</span>
                            </div>
                        </div>
                    `).join('');

                    return `
                    <div class="bg-white border-2 ${theme.border} rounded-2xl shadow-sm overflow-hidden">
                        <div class="flex items-center gap-2 px-5 py-4 ${theme.bg} border-b ${theme.border}">
                            <span class="text-2xl">${theme.icon}</span>
                            <h3 class="font-black ${theme.text} text-lg uppercase tracking-wide">${cat}</h3>
                        </div>
                        <div class="px-5 divide-y divide-slate-100">${rows}</div>
                        <div class="flex justify-between items-center px-5 py-4 mt-1 bg-slate-900 text-white">
                            <span class="font-black uppercase tracking-widest text-xs text-slate-400">Total Collected</span>
                            <span class="font-black text-2xl text-emerald-400">${catData.category_total} ብር</span>
                        </div>
                    </div>
                    `;
                }).join('');
            }
            loadSales();
        </script>
    </body>
    </html>
    """
    return inject_navigation('sales', template)

@app.get("/admin/revenue", response_class=HTMLResponse)
def admin_revenue_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Revenue Summary</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->
        
        <main class="flex-1 max-w-full px-6 py-6 grid grid-cols-1 xl:grid-cols-3 gap-6">
            <div class="xl:col-span-2 bg-white border border-slate-200 rounded-2xl p-5 shadow-sm">
                <h3 class="font-black text-slate-900 border-b border-slate-100 pb-3 mb-4 text-base">Payment Gateway Breakdown</h3>
                <div id="revenue-rows" class="divide-y divide-slate-100"></div>
            </div>
            <div class="bg-slate-900 p-6 rounded-2xl text-white flex flex-col justify-between h-fit shadow-md border border-slate-800">
                <div class="space-y-1">
                    <span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Grand Total Balance</span>
                    <h3 id="grand-total" class="text-3xl font-black text-emerald-400 tracking-tight">0 ብር</h3>
                </div>
                <button onclick="saveReportNow()" class="mt-8 text-xs font-black w-full py-3.5 bg-white/10 border border-white/20 hover:bg-white/20 text-white rounded-xl cursor-pointer transition-all active:scale-95 flex items-center justify-center gap-2">💾 Save Daily Report</button>
                <a href="/admin/report/download" class="mt-3 text-xs font-black w-full py-3.5 bg-emerald-600 hover:bg-emerald-700 text-white rounded-xl cursor-pointer transition-all active:scale-95 flex items-center justify-center gap-2 no-underline">📄 Download Daily Report</a>
                <button onclick="wipe()" class="mt-3 text-xs font-black w-full py-3.5 bg-red-950/50 border border-red-900 text-red-400 rounded-xl hover:bg-red-950/80 cursor-pointer transition-all active:scale-95">⚠️ Wipe Operational Logs</button>
            </div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->
        
        <script>
            async function loadRevenue() {
                const res = await (await fetch('/api/admin/reports/revenue')).json();
                document.getElementById('revenue-rows').innerHTML = Object.entries(res.breakdown).map(([gate, sum]) => `
                    <div class="flex justify-between items-center py-3.5 text-sm font-bold">
                        <span class="text-slate-600 flex items-center gap-2">💳 ${gate}</span>
                        <span class="text-slate-900 font-black">${sum} ብር</span>
                    </div>
                `).join('');
                document.getElementById('grand-total').innerText = res.grand_total + ' ብር';
            }
            async function saveReportNow() {
                try {
                    const res = await fetch('/api/admin/report/save', { method: 'POST' });
                    const result = await res.json();
                    if (result.status === 'success') {
                        alert('✅ Saved to Files: ' + result.filename);
                    } else {
                        alert('Could not save report. Please try again.');
                    }
                } catch (err) {
                    console.error('Save failed: ', err);
                    alert('Network error while saving.');
                }
            }
            async function wipe() {
                if(confirm("⚠️ Critical Action: This will clear out every transaction completely. A backup report will be downloaded first. Proceed?")) {
                    const res = await fetch('/api/admin/reset', { method: 'POST' });
                    const result = await res.json();
                    if (result.saved_report) {
                        window.location.href = '/admin/files/download/' + encodeURIComponent(result.saved_report);
                    }
                    loadRevenue();
                }
            }
            loadRevenue();
        </script>
    </body>
    </html>
    """
    return inject_navigation('revenue', template)

# ==================== BACKEND UTILITY API ENDPOINTS ====================

@app.get("/api/menu")
def get_menu_endpoint():
    return MENU

@app.post("/api/orders")
async def create_order_endpoint(order: OrderSubmit):
    if not order.items:
        raise HTTPException(status_code=400, detail="Cart is empty")
    total = sum(FLAT_MENU[item] * qty for item, qty in order.items.items() if item in FLAT_MENU)
    async with data_lock:
        db = load_data()
        db["last_ticket"] += 1
        tid = str(db["last_ticket"])
        db["tickets"][tid] = {"items": order.items, "total": total, "paid": False, "method": None}
        save_data(db)
    return {"status": "success", "ticket_id": tid}

@app.get("/api/admin/tickets")
async def get_tickets_endpoint():
    async with data_lock:
        db = load_data()
    return [{"id": tid, **t} for tid, t in reversed(list(db["tickets"].items()))]

@app.post("/api/admin/tickets/{ticket_id}/pay")
async def pay_ticket_endpoint(ticket_id: str, payload: PaymentSubmit):
    async with data_lock:
        db = load_data()
        if ticket_id not in db["tickets"]:
            raise HTTPException(status_code=404, detail="Ticket missing")
        db["tickets"][ticket_id]["paid"] = True
        db["tickets"][ticket_id]["method"] = payload.method
        save_data(db)
    return {"status": "success"}

@app.get("/api/admin/reports/sales")
async def get_sales_report_endpoint():
    sales = {item: 0 for item in FLAT_MENU}
    async with data_lock:
        db = load_data()
    for t in db["tickets"].values():
        if t.get("paid"):
            for item, qty in t.get("items", {}).items():
                if item in sales:
                    sales[item] += qty

    categories = {}
    grand_total_revenue = 0
    grand_total_sold = 0
    for cat, items in MENU.items():
        cat_items = [{"name": i, "sold": sales[i], "revenue": sales[i] * FLAT_MENU[i]} for i in items]
        cat_items.sort(key=lambda x: x["sold"], reverse=True)
        category_total = sum(i["revenue"] for i in cat_items)
        category_sold = sum(i["sold"] for i in cat_items)
        categories[cat] = {"items": cat_items, "category_total": category_total, "category_sold": category_sold}
        grand_total_revenue += category_total
        grand_total_sold += category_sold

    return {
        "categories": categories,
        "grand_total_revenue": grand_total_revenue,
        "grand_total_sold": grand_total_sold,
    }

@app.get("/api/admin/reports/revenue")
async def get_revenue_report_endpoint():
    async with data_lock:
        db = load_data()
    summary = {m: 0 for m in PAYMENT_METHODS}
    grand_total = 0
    for t in db["tickets"].values():
        if t.get("paid"):
            method = t.get("method", "Cash")
            summary[method] = summary.get(method, 0) + t.get("total", 0)
            grand_total += t.get("total", 0)
    return {"breakdown": summary, "grand_total": grand_total}

@app.get("/api/admin/report/preview")
async def get_report_preview_endpoint():
    async with data_lock:
        db = load_data()
    tickets_dict = db.get("tickets", {})
    tickets_sorted = sorted(tickets_dict.items(), key=lambda x: int(x[0]))

    paid = [t for _, t in tickets_sorted if t.get("paid")]
    pending = [t for _, t in tickets_sorted if not t.get("paid")]
    grand_revenue = sum(t.get("total", 0) for t in paid)

    ticket_log = [{
        "id": tid,
        "items": t.get("items", {}),
        "total": t.get("total", 0),
        "paid": t.get("paid", False),
        "method": t.get("method") if t.get("paid") else None
    } for tid, t in tickets_sorted]

    sales = {item: 0 for item in FLAT_MENU}
    for t in paid:
        for item, qty in t.get("items", {}).items():
            if item in sales:
                sales[item] += qty

    categories = {}
    for cat, cat_items in MENU.items():
        item_rows = [{"name": i, "sold": sales[i], "revenue": sales[i] * FLAT_MENU[i]} for i in cat_items]
        item_rows.sort(key=lambda x: x["sold"], reverse=True)
        categories[cat] = {"items": item_rows, "category_total": sum(i["revenue"] for i in item_rows)}

    rev_summary = {m: 0 for m in PAYMENT_METHODS}
    for t in paid:
        method = t.get("method", "Cash")
        rev_summary[method] = rev_summary.get(method, 0) + t.get("total", 0)

    return {
        "stats": {"total": len(tickets_sorted), "paid": len(paid), "pending": len(pending), "grand_total": grand_revenue},
        "ticket_log": ticket_log,
        "categories": categories,
        "revenue_breakdown": rev_summary,
        "generated_at": datetime.now().strftime('%A, %d %B %Y — %I:%M %p')
    }

@app.get("/admin/report", response_class=HTMLResponse)
def admin_report_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Smart Menu Report</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->

        <main class="flex-1 max-w-4xl px-6 py-6 space-y-6">
            <div class="flex flex-col sm:flex-row sm:items-center sm:justify-between gap-4">
                <div>
                    <h1 class="text-2xl font-black text-slate-900">📄 Smart Menu Report</h1>
                    <p class="text-xs font-bold text-slate-400" id="gen-timestamp">Loading...</p>
                </div>
                <div class="flex items-center gap-2">
                    <button onclick="saveReportNow()" class="bg-white border-2 border-slate-200 hover:bg-slate-50 active:scale-95 text-slate-700 font-black py-3 px-5 rounded-xl shadow-sm cursor-pointer transition-all text-sm uppercase tracking-wide flex items-center gap-2">💾 Save</button>
                    <a href="/admin/report/download" class="bg-emerald-600 hover:bg-emerald-700 active:scale-95 text-white font-black py-3 px-5 rounded-xl shadow-sm cursor-pointer transition-all text-sm uppercase tracking-wide no-underline flex items-center gap-2">⬇️ Download as Word Doc</a>
                </div>
            </div>

            <div id="stats-strip" class="grid grid-cols-2 sm:grid-cols-4 gap-4"></div>

            <div class="bg-white border-2 border-slate-200 rounded-2xl shadow-sm overflow-hidden">
                <div class="px-5 py-4 bg-slate-900"><h3 class="font-black text-amber-400 text-base uppercase tracking-wide">📋 Full Ticket Log</h3></div>
                <div id="ticket-log" class="divide-y divide-slate-100"></div>
            </div>

            <div>
                <h2 class="text-lg font-black text-slate-900 mb-3">📊 Sales Report</h2>
                <div id="sales-container" class="grid grid-cols-1 lg:grid-cols-2 gap-6"></div>
            </div>

            <div class="bg-white border-2 border-slate-200 rounded-2xl shadow-sm overflow-hidden">
                <div class="px-5 py-4 bg-slate-900"><h3 class="font-black text-amber-400 text-base uppercase tracking-wide">💰 Revenue Summary</h3></div>
                <div id="revenue-rows" class="divide-y divide-slate-100 px-5"></div>
            </div>

            <div class="flex gap-3">
                <button onclick="saveReportNow()" class="flex-1 bg-white border-2 border-slate-200 hover:bg-slate-50 active:scale-95 text-slate-700 font-black py-4 rounded-xl shadow-sm cursor-pointer transition-all text-sm uppercase tracking-wide">💾 Save to Files</button>
                <a href="/admin/report/download" class="flex-1 block text-center bg-emerald-600 hover:bg-emerald-700 active:scale-95 text-white font-black py-4 rounded-xl shadow-sm cursor-pointer transition-all text-sm uppercase tracking-wide no-underline">⬇️ Download as Word Doc</a>
            </div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->

        <script>
            const CATEGORY_THEME = {
                "Foods":       { icon: "🍛", bg: "bg-amber-50", text: "text-amber-700", border: "border-amber-200" },
                "Soft Drinks": { icon: "🥤", bg: "bg-sky-50",   text: "text-sky-700",   border: "border-sky-200" },
                "Hot Drinks":  { icon: "☕", bg: "bg-rose-50",  text: "text-rose-700",  border: "border-rose-200" }
            };
            const DEFAULT_THEME = { icon: "🍽️", bg: "bg-slate-50", text: "text-slate-700", border: "border-slate-200" };

            async function loadReport() {
                const data = await (await fetch('/api/admin/report/preview')).json();
                document.getElementById('gen-timestamp').innerText = "Generated: " + data.generated_at;

                document.getElementById('stats-strip').innerHTML = `
                    <div class="bg-slate-900 text-white rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Total Tickets</span><h2 class="text-2xl font-black text-emerald-400">${data.stats.total}</h2></div>
                    <div class="bg-white border-2 border-slate-200 rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Paid</span><h2 class="text-2xl font-black text-slate-900">${data.stats.paid}</h2></div>
                    <div class="bg-white border-2 border-slate-200 rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Pending</span><h2 class="text-2xl font-black text-slate-900">${data.stats.pending}</h2></div>
                    <div class="bg-slate-900 text-white rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Total Collected</span><h2 class="text-2xl font-black text-emerald-400">${data.stats.grand_total} ብር</h2></div>
                `;

                document.getElementById('ticket-log').innerHTML = data.ticket_log.length ? data.ticket_log.map(t => {
                    const itemsText = Object.entries(t.items).map(([n, q]) => `${q}x ${n}`).join(', ');
                    const statusText = t.paid ? (t.method || 'Cash') : 'Pending';
                    return `
                        <div class="flex justify-between items-center px-5 py-3 gap-4">
                            <div class="flex-1 min-w-0">
                                <span class="font-black text-slate-900">#${t.id}</span>
                                <span class="text-sm text-slate-500 ml-2">${itemsText}</span>
                            </div>
                            <span class="font-black text-emerald-600 shrink-0">${t.total} ብር</span>
                            <span class="font-black text-xs px-2 py-1 rounded-lg shrink-0 ${t.paid ? 'bg-amber-50 text-amber-700 border border-amber-200' : 'bg-red-50 text-red-500 border border-red-200'}">${statusText}</span>
                        </div>
                    `;
                }).join('') : '<div class="px-5 py-8 text-center text-slate-400 font-medium">No tickets recorded yet.</div>';

                document.getElementById('sales-container').innerHTML = Object.entries(data.categories).map(([cat, catData]) => {
                    const theme = CATEGORY_THEME[cat] || DEFAULT_THEME;
                    const rows = catData.items.map(i => `
                        <div class="flex justify-between items-center py-3 border-b border-slate-100 last:border-0">
                            <span class="text-slate-800 font-bold">${i.name}</span>
                            <div class="flex items-center gap-3">
                                <span class="font-black ${theme.bg} ${theme.text} ${theme.border} border px-2.5 py-0.5 rounded-full text-sm">x${i.sold}</span>
                                <span class="font-black text-emerald-600 min-w-[70px] text-right">${i.revenue} ብር</span>
                            </div>
                        </div>
                    `).join('');
                    return `
                    <div class="bg-white border-2 ${theme.border} rounded-2xl shadow-sm overflow-hidden">
                        <div class="flex items-center gap-2 px-5 py-3 ${theme.bg} border-b ${theme.border}">
                            <span class="text-xl">${theme.icon}</span><h3 class="font-black ${theme.text} uppercase tracking-wide text-sm">${cat}</h3>
                        </div>
                        <div class="px-5 divide-y divide-slate-100">${rows}</div>
                        <div class="flex justify-between items-center px-5 py-3 bg-slate-900 text-white">
                            <span class="font-black uppercase tracking-widest text-xs text-slate-400">Total Collected</span>
                            <span class="font-black text-xl text-emerald-400">${catData.category_total} ብር</span>
                        </div>
                    </div>
                    `;
                }).join('');

                document.getElementById('revenue-rows').innerHTML = Object.entries(data.revenue_breakdown).map(([method, amt]) => `
                    <div class="flex justify-between items-center py-3">
                        <span class="text-slate-600 font-bold">💳 ${method}</span>
                        <span class="font-black text-slate-900">${amt} ብር</span>
                    </div>
                `).join('') + `
                    <div class="flex justify-between items-center py-4 border-t-2 border-slate-900 mt-1">
                        <span class="font-black uppercase text-slate-900">Grand Total</span>
                        <span class="font-black text-2xl text-emerald-600">${data.stats.grand_total} ብር</span>
                    </div>
                `;
            }
            async function saveReportNow() {
                try {
                    const res = await fetch('/api/admin/report/save', { method: 'POST' });
                    const result = await res.json();
                    if (result.status === 'success') {
                        alert('✅ Saved to Files: ' + result.filename);
                    } else {
                        alert('Could not save report. Please try again.');
                    }
                } catch (err) {
                    console.error('Save failed: ', err);
                    alert('Network error while saving.');
                }
            }
            loadReport();
        </script>
    </body>
    </html>
    """
    return inject_navigation('report', template)


@app.get("/api/admin/files/{filename}/preview")
async def get_saved_file_preview_endpoint(filename: str):
    safe_name = os.path.basename(filename)
    path = os.path.join(REPORTS_DIR, safe_name)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="File not found")
    try:
        return parse_report_docx(path)
    except Exception as e:
        logger.error(f"Failed to parse saved report {safe_name}: {e}")
        raise HTTPException(status_code=500, detail="Could not read this file's contents")

@app.get("/admin/report/preview", response_class=HTMLResponse)
def admin_report_view_page():
    return admin_report_page()
async def download_daily_report():
    async with data_lock:
        filename = save_report_to_archive("download")
    path = os.path.join(REPORTS_DIR, filename)
    return StreamingResponse(
        open(path, "rb"),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.get("/api/admin/files")
async def list_saved_files_endpoint():
    files = []
    for fname in os.listdir(REPORTS_DIR):
        if not fname.lower().endswith(".docx"):
            continue
        fpath = os.path.join(REPORTS_DIR, fname)
        if os.path.isfile(fpath):
            stat = os.stat(fpath)
            files.append({"filename": fname, "size_kb": round(stat.st_size / 1024, 1), "modified": stat.st_mtime})
    files.sort(key=lambda x: x["modified"], reverse=True)
    return files

@app.post("/api/admin/report/save")
async def save_report_only_endpoint():
    async with data_lock:
        filename = save_report_to_archive("manual_save")
    return {"status": "success", "filename": filename}

@app.get("/api/admin/files/{filename}/preview")
async def get_file_preview_data_endpoint(filename: str):
    safe_name = os.path.basename(filename)
    json_name = os.path.splitext(safe_name)[0] + ".json"
    json_path = os.path.join(REPORTS_DIR, json_name)
    if not os.path.isfile(json_path):
        raise HTTPException(status_code=404, detail="No snapshot data available for this file")
    with open(json_path, "r") as f:
        return json.load(f)

@app.get("/admin/files/open/{filename}", response_class=HTMLResponse)
def open_saved_file_endpoint(filename: str):
    """Renders a saved report using the exact same layout as the live Smart
    Menu Report page, but loads that specific file's frozen JSON snapshot
    instead of today's live data — so each saved file shows what it actually
    contained at the time it was saved."""
    safe_name = os.path.basename(filename)
    path = os.path.join(REPORTS_DIR, safe_name)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="File not found")

    preview_url = f"/api/admin/files/{safe_name}/preview"
    download_url = f"/admin/files/download/{safe_name}"

    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Saved Report</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->

        <main class="flex-1 max-w-4xl px-6 py-6 space-y-6">
            <div class="flex flex-col sm:flex-row sm:items-center sm:justify-between gap-4">
                <div>
                    <h1 class="text-2xl font-black text-slate-900">📄 Saved Report</h1>
                    <p class="text-xs font-bold text-slate-400" id="gen-timestamp">Loading...</p>
                    <p class="text-[11px] font-bold text-amber-600 uppercase tracking-wide mt-1">""" + safe_name + """</p>
                </div>
                <div class="flex items-center gap-2">
                    <a href='""" + download_url + """' class="bg-emerald-600 hover:bg-emerald-700 active:scale-95 text-white font-black py-3 px-5 rounded-xl shadow-sm cursor-pointer transition-all text-sm uppercase tracking-wide no-underline flex items-center gap-2">⬇️ Download as Word Doc</a>
                </div>
            </div>

            <div id="stats-strip" class="grid grid-cols-2 sm:grid-cols-4 gap-4"></div>

            <div class="bg-white border-2 border-slate-200 rounded-2xl shadow-sm overflow-hidden">
                <div class="px-5 py-4 bg-slate-900"><h3 class="font-black text-amber-400 text-base uppercase tracking-wide">📋 Full Ticket Log</h3></div>
                <div id="ticket-log" class="divide-y divide-slate-100"></div>
            </div>

            <div>
                <h2 class="text-lg font-black text-slate-900 mb-3">📊 Sales Report</h2>
                <div id="sales-container" class="grid grid-cols-1 lg:grid-cols-2 gap-6"></div>
            </div>

            <div class="bg-white border-2 border-slate-200 rounded-2xl shadow-sm overflow-hidden">
                <div class="px-5 py-4 bg-slate-900"><h3 class="font-black text-amber-400 text-base uppercase tracking-wide">💰 Revenue Summary</h3></div>
                <div id="revenue-rows" class="divide-y divide-slate-100 px-5"></div>
            </div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->

        <script>
            const CATEGORY_THEME = {
                "Foods":       { icon: "🍛", bg: "bg-amber-50", text: "text-amber-700", border: "border-amber-200" },
                "Soft Drinks": { icon: "🥤", bg: "bg-sky-50",   text: "text-sky-700",   border: "border-sky-200" },
                "Hot Drinks":  { icon: "☕", bg: "bg-rose-50",  text: "text-rose-700",  border: "border-rose-200" }
            };
            const DEFAULT_THEME = { icon: "🍽️", bg: "bg-slate-50", text: "text-slate-700", border: "border-slate-200" };

            async function loadReport() {
                const res = await fetch('""" + preview_url + """');
                if (!res.ok) {
                    document.getElementById('gen-timestamp').innerText = "No saved snapshot data found for this file.";
                    return;
                }
                const data = await res.json();
                document.getElementById('gen-timestamp').innerText = "Generated: " + data.generated_at;

                document.getElementById('stats-strip').innerHTML = `
                    <div class="bg-slate-900 text-white rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Total Tickets</span><h2 class="text-2xl font-black text-emerald-400">${data.stats.total}</h2></div>
                    <div class="bg-white border-2 border-slate-200 rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Paid</span><h2 class="text-2xl font-black text-slate-900">${data.stats.paid}</h2></div>
                    <div class="bg-white border-2 border-slate-200 rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Pending</span><h2 class="text-2xl font-black text-slate-900">${data.stats.pending}</h2></div>
                    <div class="bg-slate-900 text-white rounded-2xl p-4 text-center"><span class="text-[10px] uppercase tracking-widest font-black text-slate-400 block">Total Collected</span><h2 class="text-2xl font-black text-emerald-400">${data.stats.grand_total} ብር</h2></div>
                `;

                document.getElementById('ticket-log').innerHTML = data.ticket_log.length ? data.ticket_log.map(t => {
                    const itemsText = Object.entries(t.items).map(([n, q]) => `${q}x ${n}`).join(', ');
                    const statusText = t.paid ? (t.method || 'Cash') : 'Pending';
                    return `
                        <div class="flex justify-between items-center px-5 py-3 gap-4">
                            <div class="flex-1 min-w-0">
                                <span class="font-black text-slate-900">#${t.id}</span>
                                <span class="text-sm text-slate-500 ml-2">${itemsText}</span>
                            </div>
                            <span class="font-black text-emerald-600 shrink-0">${t.total} ብር</span>
                            <span class="font-black text-xs px-2 py-1 rounded-lg shrink-0 ${t.paid ? 'bg-amber-50 text-amber-700 border border-amber-200' : 'bg-red-50 text-red-500 border border-red-200'}">${statusText}</span>
                        </div>
                    `;
                }).join('') : '<div class="px-5 py-8 text-center text-slate-400 font-medium">No tickets recorded.</div>';

                document.getElementById('sales-container').innerHTML = Object.entries(data.categories).map(([cat, catData]) => {
                    const theme = CATEGORY_THEME[cat] || DEFAULT_THEME;
                    const rows = catData.items.map(i => `
                        <div class="flex justify-between items-center py-3 border-b border-slate-100 last:border-0">
                            <span class="text-slate-800 font-bold">${i.name}</span>
                            <div class="flex items-center gap-3">
                                <span class="font-black ${theme.bg} ${theme.text} ${theme.border} border px-2.5 py-0.5 rounded-full text-sm">x${i.sold}</span>
                                <span class="font-black text-emerald-600 min-w-[70px] text-right">${i.revenue} ብር</span>
                            </div>
                        </div>
                    `).join('');
                    return `
                    <div class="bg-white border-2 ${theme.border} rounded-2xl shadow-sm overflow-hidden">
                        <div class="flex items-center gap-2 px-5 py-3 ${theme.bg} border-b ${theme.border}">
                            <span class="text-xl">${theme.icon}</span><h3 class="font-black ${theme.text} uppercase tracking-wide text-sm">${cat}</h3>
                        </div>
                        <div class="px-5 divide-y divide-slate-100">${rows}</div>
                        <div class="flex justify-between items-center px-5 py-3 bg-slate-900 text-white">
                            <span class="font-black uppercase tracking-widest text-xs text-slate-400">Total Collected</span>
                            <span class="font-black text-xl text-emerald-400">${catData.category_total} ብር</span>
                        </div>
                    </div>
                    `;
                }).join('');

                document.getElementById('revenue-rows').innerHTML = Object.entries(data.revenue_breakdown).map(([method, amt]) => `
                    <div class="flex justify-between items-center py-3">
                        <span class="text-slate-600 font-bold">💳 ${method}</span>
                        <span class="font-black text-slate-900">${amt} ብር</span>
                    </div>
                `).join('') + `
                    <div class="flex justify-between items-center py-4 border-t-2 border-slate-900 mt-1">
                        <span class="font-black uppercase text-slate-900">Grand Total</span>
                        <span class="font-black text-2xl text-emerald-600">${data.stats.grand_total} ብር</span>
                    </div>
                `;
            }
            loadReport();
        </script>
    </body>
    </html>
    """
    return inject_navigation('files', template)

@app.get("/admin/files/download/{filename}")
async def download_saved_file_endpoint(filename: str):
    safe_name = os.path.basename(filename)
    path = os.path.join(REPORTS_DIR, safe_name)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="File not found")
    return StreamingResponse(
        open(path, "rb"),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'}
    )

@app.get("/admin/files", response_class=HTMLResponse)
def admin_files_page():
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
        <title>Files</title><link rel="stylesheet" href="/static/tailwind.css">
    </head>
    <body class="bg-slate-50 text-slate-800 antialiased min-h-screen">
        <!-- START_WRAPPER_PLACEHOLDER -->
        <!-- NAV_SIDEBAR_PLACEHOLDER -->
        <main class="flex-1 max-w-4xl px-6 py-6 space-y-6">
            <div>
                <h1 class="text-2xl font-black text-slate-900">📁 Saved Files</h1>
                <p class="text-xs font-bold text-slate-400 uppercase tracking-widest">Every report ever generated, kept here for re-download</p>
            </div>
            <div id="files-list" class="space-y-3"></div>
        </main>
        <!-- END_WRAPPER_PLACEHOLDER -->
        <script>
            async function loadFiles() {
                const files = await (await fetch('/api/admin/files')).json();
                const list = document.getElementById('files-list');
                if (files.length === 0) {
                    list.innerHTML = '<div class="text-center py-16 bg-white border-2 border-dashed border-slate-200 rounded-2xl text-slate-400 font-medium">No saved reports yet. Download one from the Smart Menu Report page.</div>';
                    return;
                }
                list.innerHTML = files.map(f => `
                    <div class="flex items-center justify-between bg-white border-2 border-slate-200 rounded-2xl px-5 py-4 shadow-sm">
                        <div class="flex items-center gap-3 min-w-0">
                            <span class="text-2xl shrink-0">📄</span>
                            <div class="min-w-0">
                                <p class="font-black text-slate-900 truncate">${f.filename}</p>
                                <p class="text-xs text-slate-400 font-bold">${f.size_kb} KB • ${new Date(f.modified * 1000).toLocaleString()}</p>
                            </div>
                        </div>
                        <div class="flex items-center gap-2 shrink-0">
                            <a href="/admin/files/open/${encodeURIComponent(f.filename)}" target="_blank" class="bg-white border-2 border-slate-200 hover:bg-slate-50 active:scale-95 text-slate-700 font-black py-2.5 px-4 rounded-xl text-xs uppercase tracking-wide no-underline transition-all">👁️ Open</a>
                            <a href="/admin/files/download/${encodeURIComponent(f.filename)}" class="bg-emerald-600 hover:bg-emerald-700 active:scale-95 text-white font-black py-2.5 px-4 rounded-xl text-xs uppercase tracking-wide no-underline transition-all">⬇️ Download</a>
                        </div>
                    </div>
                `).join('');
            }
            loadFiles();
        </script>
    </body>
    </html>
    """
    return inject_navigation('files', template)

@app.post("/api/admin/reset")
async def reset_data_endpoint():
    async with data_lock:
        filename = save_report_to_archive("auto_wipe")
        save_data({"last_ticket": 0, "tickets": {}})
    return {"status": "success", "saved_report": filename}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=True)